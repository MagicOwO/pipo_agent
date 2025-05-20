import os
from docx import Document
import pyglove as pg
import langfun as lf
import sys
from typing import Dict, List, Literal

def read_word_document(file_path: str) -> str:
    """
    Read content from a Word document and return the text.
    
    Args:
        file_path (str): Path to the Word document
        
    Returns:
        str: Text content of the document
    """
    try:
        # Load the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    
    except Exception as e:
        print(f"Error reading the document: {str(e)}")
        return None

class AudienceFilter(pg.Object):
    operator: Literal["lt", "not_in"]
    value: str

class MessageTask(pg.Object):
    day_index: int
    send_time: str
    audience_filter: List[AudienceFilter]
    message_type: Literal["text", "image", "video", "voice"]
    message_content: int

class StructuredContent(pg.Object):
    message_tasks: List[MessageTask]


def main():
    instruction_file = "D:/Dev/download/说明文档-量子POC.docx"
    input_file = "D:/Dev/download/输入文档.docx"
    instruction = read_word_document(instruction_file)
    input_content = read_word_document(input_file)
    prompt = f"instruction: {instruction}\n\ninput_content: {input_content}"
    structured_content = lf.query(
        prompt,
        StructuredContent,
        lm=lf.llms.Gpt4o(api_key=os.getenv("OPENAI_API_KEY"))
    )
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    
    output_file = os.path.join(output_dir, "structured_content.txt")
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(str(structured_content))
    print(f"Structured content written to {output_file}")

if __name__ == "__main__":
    main()
